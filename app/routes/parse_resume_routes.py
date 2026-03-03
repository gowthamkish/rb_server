"""
Resume parsing route.

POST /api/parse-resume
    Accepts a .pdf, .docx, or .doc file, extracts text, and returns structured
    resume data (personalInfo, experiences, education, skills).
"""

import io
import os
import subprocess
import tempfile
import traceback

import pdfplumber
import pymupdf  # PyMuPDF – most robust PDF parser
from docx import Document
from fastapi import APIRouter, HTTPException, UploadFile
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pypdf import PdfReader

from app.utils.resume_parser import parse_resume_text

router = APIRouter()

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def _extract_text_from_pdf(buffer: bytes) -> str:
    """
    Extract plain text from a PDF buffer.
    Tries multiple strategies in order of robustness:
      1. PyMuPDF with text mode         – best for most PDFs
      2. PyMuPDF with words mode        – handles some edge cases
      3. pdfplumber with word fallback  – layout-aware
      4. pdfminer.six                   – handles some encodings others miss
    Returns the first non-empty result.
    """
    # Strategy 1 & 2: PyMuPDF (handles compressed/encoded fonts best)
    try:
        doc = pymupdf.open(stream=buffer, filetype="pdf")
        # Attempt to open without password if permission-locked
        if doc.is_encrypted:
            doc.authenticate("")  # try empty password

        text_parts: list[str] = []
        for page in doc:
            t = page.get_text("text")  # type: ignore[attr-defined]
            if not t.strip():
                # fallback: extract word-by-word within pymupdf
                words = page.get_text("words")  # type: ignore[attr-defined]
                if words:
                    t = " ".join(w[4] for w in words)
            if t:
                text_parts.append(t)
        doc.close()
        text = "\n".join(text_parts)
        if text.strip():
            print(f"[parse-resume] PyMuPDF extracted {len(text.strip())} chars")
            return text
    except Exception:
        traceback.print_exc()

    # Strategy 3: pdfplumber with relaxed tolerances + word fallback
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(buffer)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if not page_text:
                    words = page.extract_words()
                    page_text = " ".join(w["text"] for w in words) if words else ""
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts)
        if text.strip():
            print(f"[parse-resume] pdfplumber extracted {len(text.strip())} chars")
            return text
    except Exception:
        traceback.print_exc()

    # Strategy 4: pdfminer.six
    try:
        text = pdfminer_extract_text(io.BytesIO(buffer))
        if text.strip():
            print(f"[parse-resume] pdfminer extracted {len(text.strip())} chars")
            return text
    except Exception:
        traceback.print_exc()

    # Strategy 5: pypdf (different font/encoding handling)
    try:
        reader = PdfReader(io.BytesIO(buffer))
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        text = "\n".join(parts)
        if text.strip():
            print(f"[parse-resume] pypdf extracted {len(text.strip())} chars")
            return text
    except Exception:
        traceback.print_exc()

    return ""


def _extract_text_from_docx(buffer: bytes) -> str:
    """Extract plain text from a DOCX buffer using python-docx."""
    doc = Document(io.BytesIO(buffer))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)
    return "\n".join(paragraphs)


def _extract_text_from_doc(buffer: bytes) -> str:
    """
    Extract plain text from a legacy .doc (Word 97-2003) buffer using antiword.
    Falls back to reading the raw buffer bytes for any readable ASCII text
    if antiword is unavailable.
    """
    tmp_path = ""
    try:
        # Write buffer to a temp file (antiword requires a file path)
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(buffer)
            tmp_path = tmp.name

        result = subprocess.run(
            ["antiword", "-w", "0", tmp_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            print(f"[parse-resume] antiword extracted {len(result.stdout.strip())} chars")
            return result.stdout
        if result.stderr:
            print(f"[parse-resume] antiword stderr: {result.stderr.strip()}")
    except FileNotFoundError:
        print("[parse-resume] antiword not found on PATH")
    except subprocess.TimeoutExpired:
        print("[parse-resume] antiword timed out")
    except Exception:
        traceback.print_exc()
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # Fallback: naive byte scan for printable ASCII (last resort)
    try:
        text = buffer.decode("latin-1", errors="ignore")
        printable = "".join(c if c.isprintable() or c in "\n\t" else " " for c in text)
        # Collapse long whitespace runs that are typical in binary .doc files
        import re
        printable = re.sub(r"[ \t]{3,}", "  ", printable)
        printable = re.sub(r"\n{3,}", "\n\n", printable)
        if len(printable.strip()) > 100:
            print(f"[parse-resume] fallback byte scan extracted {len(printable.strip())} chars")
            return printable
    except Exception:
        pass

    return ""


@router.post("/")
async def parse_resume(file: UploadFile):
    """
    Upload a .pdf, .docx, or .doc resume file and receive structured JSON
    with personalInfo, experiences, education, and skills.
    """
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded")

    filename: str = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in {"pdf", "docx", "doc"}:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Allowed: .pdf, .docx, .doc",
        )

    buffer = await file.read()
    print(f"[parse-resume] Received file: {filename} ({len(buffer)} bytes, ext={ext})")
    if len(buffer) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 10 MB)")

    if len(buffer) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty (0 bytes)")

    # Extract raw text
    try:
        if ext == "pdf":
            raw_text = _extract_text_from_pdf(buffer)
        elif ext == "doc":
            raw_text = _extract_text_from_doc(buffer)
        else:
            raw_text = _extract_text_from_docx(buffer)
    except Exception as exc:
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from file: {exc}",
        )

    print(f"[parse-resume] Extracted text length: {len(raw_text.strip())} chars")

    if not raw_text.strip():
        raise HTTPException(
            status_code=422,
            detail=(
                "Could not extract any text from the uploaded file. "
                "For PDFs, the file may be image-based or use non-standard fonts. "
                "For .doc files, please resave as .docx in Microsoft Word and re-upload."
            ),
        )

    # Parse extracted text into structured resume data
    try:
        result = parse_resume_text(raw_text)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse resume content: {exc}",
        )

    return result
